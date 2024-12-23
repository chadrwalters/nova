#!/usr/bin/env python3

import os
import magic
import logging
import json
from pathlib import Path
from typing import Optional, Dict, Any
from bs4 import BeautifulSoup
import html2text
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image, UnidentifiedImageError
import pillow_heif
import chardet

# Register HEIF opener with Pillow
pillow_heif.register_heif_opener()

class AttachmentProcessor:
    """Base class for attachment processors."""
    
    def __init__(self, file_path: Path, output_dir: Path):
        self.file_path = file_path
        self.output_dir = output_dir
        self.mime_type = magic.from_file(str(file_path), mime=True)
        self.metadata: Dict[str, Any] = {}
        
        # Set up logging
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(logging.DEBUG)
        
        # Create file handler
        log_file = output_dir / "attachment_processing.log"
        fh = logging.FileHandler(str(log_file))
        fh.setLevel(logging.DEBUG)
        
        # Create formatter
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        fh.setFormatter(formatter)
        
        # Add handler
        self.logger.addHandler(fh)
    
    def process(self) -> Optional[str]:
        """Process the attachment and return markdown content."""
        try:
            return self._process()
        except Exception as e:
            self.logger.error(f"Error processing {self.file_path}: {str(e)}")
            return None
    
    def _process(self) -> Optional[str]:
        """Internal processing method to be implemented by subclasses."""
        raise NotImplementedError
    
    def _save_metadata(self):
        """Save metadata to a JSON file."""
        metadata_file = self.output_dir / f"{self.file_path.stem}_metadata.json"
        with open(metadata_file, 'w') as f:
            json.dump(self.metadata, f, indent=2)

class HTMLProcessor(AttachmentProcessor):
    """Process HTML files."""
    
    def _process(self) -> Optional[str]:
        with open(self.file_path, 'rb') as f:
            raw_content = f.read()
            result = chardet.detect(raw_content)
            content = raw_content.decode(result['encoding'] or 'utf-8')
        
        # Parse HTML
        soup = BeautifulSoup(content, 'html.parser')
        
        # Extract metadata
        self.metadata['title'] = soup.title.string if soup.title else None
        
        # Convert to markdown
        h = html2text.HTML2Text()
        h.body_width = 0  # No wrapping
        markdown = h.handle(content)
        
        return markdown

class CSVProcessor(AttachmentProcessor):
    """Process CSV and Excel files."""
    
    def _process(self) -> Optional[str]:
        try:
            # Try reading as Excel first
            if self.file_path.suffix.lower() in ['.xlsx', '.xls']:
                df = pd.read_excel(self.file_path)
            else:
                # For CSV, detect encoding
                with open(self.file_path, 'rb') as f:
                    raw_content = f.read()
                    result = chardet.detect(raw_content)
                    encoding = result['encoding'] or 'utf-8'
                
                df = pd.read_csv(self.file_path, encoding=encoding)
            
            # Convert to markdown table
            markdown = df.to_markdown(index=False)
            
            # Store metadata
            self.metadata['columns'] = df.columns.tolist()
            self.metadata['rows'] = len(df)
            
            return markdown
        except Exception as e:
            self.logger.error(f"Error processing CSV/Excel file: {str(e)}")
            return None

class PDFProcessor(AttachmentProcessor):
    """Process PDF files."""
    
    def _process(self) -> Optional[str]:
        try:
            reader = PdfReader(self.file_path)
            
            # Extract text from all pages
            text = []
            for page in reader.pages:
                text.append(page.extract_text())
            
            # Store metadata
            self.metadata['pages'] = len(reader.pages)
            self.metadata['title'] = reader.metadata.get('/Title', '')
            self.metadata['author'] = reader.metadata.get('/Author', '')
            
            return "\n\n".join(text)
        except Exception as e:
            self.logger.error(f"Error processing PDF file: {str(e)}")
            return None

class ImageProcessor(AttachmentProcessor):
    """Process image files."""
    
    def _process(self) -> Optional[str]:
        try:
            # Convert HEIC to JPEG if needed
            if self.file_path.suffix.lower() in ['.heic', '.heif']:
                self.logger.info(f"Converting HEIC image: {self.file_path}")
                jpeg_path = self.output_dir / f"{self.file_path.stem}.jpg"
                
                # Open HEIC image using pillow-heif
                heif_file = pillow_heif.read_heif(str(self.file_path))
                image = Image.frombytes(
                    heif_file.mode,
                    heif_file.size,
                    heif_file.data,
                    "raw",
                    heif_file.mode,
                    heif_file.stride,
                )
                
                # Save as JPEG
                image.save(jpeg_path, 'JPEG', quality=85)
                self.logger.info(f"Converted HEIC to JPEG: {jpeg_path}")
                
                # Update file path to use the converted image
                self.file_path = jpeg_path
            
            # Open and process image
            with Image.open(self.file_path) as img:
                # Store metadata
                self.metadata['format'] = img.format
                self.metadata['mode'] = img.mode
                self.metadata['size'] = img.size
                
                # Generate markdown image reference
                return f"![{self.file_path.stem}]({self.file_path})"
        except UnidentifiedImageError:
            self.logger.error(f"Unidentified image format: {self.file_path}")
            return None
        except Exception as e:
            self.logger.error(f"Error processing image file: {str(e)}")
            return None

class OfficeProcessor(AttachmentProcessor):
    """Process Microsoft Office documents."""
    
    def _process(self) -> Optional[str]:
        try:
            doc = Document(self.file_path)
            
            # Extract text from paragraphs
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                df = [[cell.text for cell in row.cells] for row in table.rows]
                headers = df[0]
                data = df[1:]
                table_df = pd.DataFrame(data, columns=headers)
                text.append(table_df.to_markdown(index=False))
            
            # Store metadata
            self.metadata['paragraphs'] = len(doc.paragraphs)
            self.metadata['tables'] = len(doc.tables)
            
            return "\n\n".join(text)
        except Exception as e:
            self.logger.error(f"Error processing Office document: {str(e)}")
            return None

def create_processor(file_path: Path, output_dir: Path) -> Optional[AttachmentProcessor]:
    """Create an appropriate processor based on file type."""
    
    suffix = file_path.suffix.lower()
    mime_type = magic.from_file(str(file_path), mime=True)
    
    if suffix in ['.html', '.htm'] or mime_type == 'text/html':
        return HTMLProcessor(file_path, output_dir)
    elif suffix in ['.csv', '.xlsx', '.xls']:
        return CSVProcessor(file_path, output_dir)
    elif suffix == '.pdf' or mime_type == 'application/pdf':
        return PDFProcessor(file_path, output_dir)
    elif suffix in ['.jpg', '.jpeg', '.png', '.gif', '.heic', '.heif']:
        return ImageProcessor(file_path, output_dir)
    elif suffix in ['.doc', '.docx']:
        return OfficeProcessor(file_path, output_dir)
    
    return None

if __name__ == "__main__":
    # Example usage
    input_file = Path("example.docx")
    output_dir = Path("output")
    
    processor = create_processor(input_file, output_dir)
    if processor:
        markdown = processor.process()
        if markdown:
            print("Successfully processed file") 