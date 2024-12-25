"""Core attachment processing functionality."""

import os
import magic
import logging
import json
import chardet
from pathlib import Path
from typing import Optional, Dict, Any, List
from bs4 import BeautifulSoup
import html2text
import pandas as pd
import markitdown
from docx import Document
from PIL import Image, UnidentifiedImageError
import tempfile
import subprocess

from ..base_handler import BaseHandler
from nova.core.image_processor import ImageProcessor

class AttachmentHandler(BaseHandler):
    """Base handler for processing file attachments."""
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """Initialize the attachment handler.
        
        Args:
            config: Optional configuration overrides
        """
        super().__init__(config)
        
        # Set up logging
        self.logger = logging.getLogger(__name__)
        self.logger.setLevel(logging.DEBUG)
        
        # Initialize image processor
        self.image_processor = ImageProcessor(config=config)
        
        # Configure supported types
        self.supported_types = {
            'text/html': self._process_html,
            'text/csv': self._process_csv,
            'application/pdf': self._process_pdf,
            'image/': self._process_image,
            'application/vnd.openxmlformats-officedocument': self._process_office
        }
    
    def can_handle(self, file_path: Path, attachments: Optional[List[Path]] = None) -> bool:
        """Check if file can be processed.
        
        Args:
            file_path: Path to the file to check
            attachments: Optional list of attachments (not used)
            
        Returns:
            bool: True if file type is supported
        """
        mime_type = magic.from_file(str(file_path), mime=True)
        
        # Check if it's an image
        if mime_type.startswith('image/'):
            return file_path.suffix.lower() in self.image_processor.supported_formats
        
        # Check other types
        return any(mime_type.startswith(t) for t in self.supported_types.keys())
    
    async def process(
        self, 
        file_path: Path, 
        context: Dict[str, Any],
        attachments: Optional[List[Path]] = None
    ) -> Dict[str, Any]:
        """Process an attachment file.
        
        Args:
            file_path: Path to the file to process
            context: Processing context
            attachments: Optional list of attachments (not used)
            
        Returns:
            Dict containing:
                - content: Processed content as markdown
                - metadata: File metadata
                - errors: Processing errors
                - processed_attachments: Empty list
                - performance: Performance metrics (for image processing)
        """
        result = {
            'content': '',
            'metadata': {},
            'errors': [],
            'processed_attachments': []
        }
        
        try:
            # Get MIME type
            mime_type = magic.from_file(str(file_path), mime=True)
            
            # Find appropriate processor
            processor = None
            for type_prefix, proc_func in self.supported_types.items():
                if mime_type.startswith(type_prefix):
                    processor = proc_func
                    break
            
            if processor:
                # Process the file
                content = await processor(file_path)
                if content:
                    result['content'] = content
                    
                # Add basic metadata
                result['metadata'].update({
                    'mime_type': mime_type,
                    'size': os.path.getsize(file_path),
                    'filename': file_path.name
                })
                
                # Add performance metrics for image processing
                if mime_type.startswith('image/'):
                    result['performance'] = self.image_processor.get_performance_report()
            else:
                result['errors'].append(f"No processor found for MIME type: {mime_type}")
            
        except Exception as e:
            result['errors'].append(str(e))
        
        return result
    
    def validate_output(self, result: Dict[str, Any]) -> bool:
        """Validate processing results.
        
        Args:
            result: Processing results to validate
            
        Returns:
            bool: True if results are valid
        """
        required_keys = {'content', 'metadata', 'errors', 'processed_attachments'}
        return (
            all(key in result for key in required_keys) and
            isinstance(result['content'], str) and
            isinstance(result['metadata'], dict) and
            isinstance(result['errors'], list) and
            isinstance(result['processed_attachments'], list)
        )
    
    async def _process_html(self, file_path: Path) -> Optional[str]:
        """Process HTML files."""
        with open(file_path, 'rb') as f:
            raw_content = f.read()
            result = chardet.detect(raw_content)
            content = raw_content.decode(result['encoding'] or 'utf-8')
        
        # Parse HTML
        soup = BeautifulSoup(content, 'html.parser')
        
        # Convert to markdown
        h = html2text.HTML2Text()
        h.body_width = 0  # No wrapping
        return h.handle(content)
    
    async def _process_csv(self, file_path: Path) -> Optional[str]:
        """Process CSV and Excel files."""
        try:
            # Try reading as Excel first
            if file_path.suffix.lower() in ['.xlsx', '.xls']:
                df = pd.read_excel(file_path)
            else:
                # For CSV, detect encoding
                with open(file_path, 'rb') as f:
                    raw_content = f.read()
                    result = chardet.detect(raw_content)
                    encoding = result['encoding'] or 'utf-8'
                
                df = pd.read_csv(file_path, encoding=encoding)
            
            # Convert to markdown table
            return df.to_markdown(index=False)
            
        except Exception as e:
            self.logger.error(f"Error processing CSV/Excel file: {str(e)}")
            return None
    
    async def _process_pdf(self, file_path: Path) -> Optional[str]:
        """Process PDF files."""
        try:
            # Convert PDF to markdown using markitdown
            converter = markitdown.MarkItDown()
            result = converter.convert_local(str(file_path), output_format="markdown")
            return result.text_content
            
        except Exception as e:
            self.logger.error(f"Error processing PDF file: {str(e)}")
            return None
    
    async def _process_image(self, file_path: Path) -> Optional[str]:
        """Process image files using ImageProcessor."""
        try:
            # Process image
            process_result = await self.image_processor.process_image(file_path)
            
            # Convert HEIC to JPEG if needed
            if file_path.suffix.lower() in ['.heic', '.heif']:
                self.logger.info(f"Converting HEIC image: {file_path}")
                jpeg_path = await self.image_processor.convert_heic_to_jpg(file_path)
                file_path = jpeg_path
            
            # Optimize image if configured
            if self.config.get('optimize_quality', False):
                preset = self.config.get('optimization_preset', 'web')
                file_path = await self.image_processor.optimize_image(file_path, preset)
            
            # Generate markdown image reference with metadata
            metadata = await self.image_processor.extract_metadata(file_path)
            dimensions = metadata['basic']['dimensions']
            alt_text = f"{file_path.stem} ({dimensions[0]}x{dimensions[1]})"
            
            return f"![{alt_text}]({file_path})"
            
        except Exception as e:
            self.logger.error(f"Error processing image file: {str(e)}")
            return None
    
    async def _process_office(self, file_path: Path) -> Optional[str]:
        """Process Microsoft Office documents."""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text = []
            for para in doc.paragraphs:
                text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                df = [[cell.text for cell in row.cells] for row in table.rows]
                headers = df[0]
                data = df[1:]
                table_df = pd.DataFrame(data, columns=headers)
                text.append(table_df.to_markdown(index=False))
            
            return "\n\n".join(text)
            
        except Exception as e:
            self.logger.error(f"Error processing Office document: {str(e)}")
            return None
    
    async def cleanup(self) -> None:
        """Clean up resources."""
        # Add any cleanup needed
        pass 